# coding=utf-8
import docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
# from DBM.mongo_handle import MongoHandle
import pandas as pd
# from global_config import *
import datetime
from jqdatasdk import *
from configDB import *
import matplotlib.pyplot as plt
auth(JOINQUANT_USER, JOINQUANT_PW)


def make_doc(self, today=None):
    """
    生成报表
    :return:
    """
    today = datetime.date.today() if not today else pd.to_datetime(today)
    # last_trade_date = basic_data.find_tradeday(-1, today)
    # bfr_monday_date = basic_data.find_tradeday(-1, today + datetime.timedelta(days=-today.weekday(), weeks=0))
    # bfr_month_date = basic_data.find_tradeday(-1, today.replace(day=1))
    # bfr_year_date = basic_data.find_tradeday(-1, today.replace(month=1, day=1))
    last_trade_date = today
    bfr_monday_date = today
    bfr_month_date = today
    bfr_year_date = today
    today = today.strftime('%Y%m%d')
    fund_list = list()
    chat_paths = list()
    for name, obj in self.hosts.items():
        chat_path = obj.draw_chat(today)  # 画图
        fund = obj.fund_now.copy()  # 资金
        fund_his = obj.fund_his.copy().append(fund)
        fund = fund_his.loc[fund_his['交易日'] == today]  # 资金
        key = fund_his['交易日'] >= last_trade_date
        fund['当日收益'] = f"{round((fund.iloc[-1]['净值'] / fund_his.loc[key, '净值'].iloc[0] - 1) * 100, 2)}%"
        key = fund_his['交易日'] >= bfr_monday_date
        fund['当周收益'] = f"{round((fund.iloc[-1]['净值'] / fund_his.loc[key, '净值'].iloc[0] - 1) * 100 ,2)}%"
        key = fund_his['交易日'] >= bfr_month_date
        fund['当月收益'] = f"{round((fund.iloc[-1]['净值'] / fund_his.loc[key, '净值'].iloc[0] - 1) * 100, 2)}%"
        key = fund_his['交易日'] >= bfr_year_date
        fund['当年收益'] = f"{round((fund.iloc[-1]['净值'] / fund_his.loc[key, '净值'].iloc[0] - 1) * 100, 2)}%"
        fund = fund[['策略名称', '资产总值', '证券市值', '总盈亏', '净值份额', '净值', '当日收益', '当周收益',
                     '当月收益', '当年收益']]
        fund[['资产总值', '证券市值', '总盈亏']] = round(fund.iloc[-1]['资产总值'], 2), \
                                        round(fund.iloc[-1]['证券市值'], 2), round(fund.iloc[-1]['总盈亏'], 2)
        fund['净值'] = round(fund.iloc[-1]['净值'], 4)
        fund_list.append(fund)
        chat_paths.append(chat_path)
    funds = pd.concat(fund_list)
    funds.sort_values('策略名称', inplace=True)
    document = Document()
    # 添加标题，格式2
    for section in document.sections:
        section.page_width = Inches(10)
    document.add_heading()


def save_df_to_doc(document, test_df, word=None):
    """
    将结果按照dataframe的形式存入doc文件
    :param document: 存入的文档类
    :param test_df: 需要保存的df
    :return:

    """
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # add_paragraph表示添加一个段落
    if word:
        document.add_paragraph(u'\n%s' % (word))
    if len(test_df.columns) == 0:
        # document.add_paragraph(u'\n%s' % ('无'))
        return

    # 添加一个表格--行数和列数，行数多加一行，需要将列名同时保存
    t = document.add_table(test_df.shape[0] + 1, test_df.shape[1], style="Table Grid")
    t.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格整体居中
    # 将每列列名保存到表格中
    for j in range(test_df.shape[-1]):
        t.cell(0, j).text = test_df.columns[j]
        t.cell(0, j).width = Inches(1.85)

    # 将每列数据保存到新建的表格中
    for i in range(test_df.shape[0]):
        for j in range(test_df.shape[-1]):
            # 第一行保存的是列名，所以数据保存时，行数要加1
            t.cell(i + 1, j).text = str(test_df.values[i, j])

    for row in t.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(7.5)
    for col in range(test_df.shape[1]):
        t.cell(0, col).width = Inches(1.65)


def get_date(calen, today):
    next_tradeday = get_trade_days(start_date=today + datetime.timedelta(days=1), end_date='2030-01-01')[0]
    if datetime.datetime.now().hour >= 15:
        calen.append(next_tradeday)
    EndDate = calen[-1]
    StartDate = calen[0]
    hq_last_date = calen[-2]
    return calen, next_tradeday, EndDate, StartDate, str(hq_last_date)[:10]


if __name__ == '__main__':
    date = '20200710'
    path = 'G:/trading/trading_report/'
    account1 = '21900576'
    account2 = '85030120'
    bars = 5
    today = datetime.date.today()
    calen = get_trade_days(count=bars)
    calen = list(calen)
    calen, next_tradeday, EndDate, StartDate, hq_last_date = get_date(calen, today)

    hq_last_date = hq_last_date[:4] + hq_last_date[5:7] + hq_last_date[8:]
    today = datetime.date.strftime(today, '%Y%m%d')
    document = Document()
    document.add_heading(f'交易报告{today}', level=0)
    document.add_heading(f'账户汇总')
    # document.add_picture(f'{path}/account_{date}.png', width=Inches(6.0))
    # document.add_picture(f'{path}/JZ_ZG_{date}.png', width=Inches(6.0))
    # document.add_picture(f'{path}/JZ_ROE_{date}.png', width=Inches(6.0))

    fund_df1 = pd.read_excel(path + 'fund_' + account1 + '_' + today + '.xlsx')\
        .assign(交易日=lambda df: df['交易日'].apply(lambda x: str(int(x+0.1))))

    fund_df2 = pd.read_excel(path + 'fund_' + account2 + '_' + today + '.xlsx')\
        .assign(交易日=lambda df: df['交易日'].apply(lambda x: str(int(x+0.1))))

    fund_today1 = fund_df1[fund_df1['交易日'] == today].loc[:, ['交易日', '动态权益', '保证金']]\
        .assign(账户='浙商期货')\
        .assign(初始资产=4000000)\
        .assign(动态权益=lambda df: df['动态权益'] - 1)\
        .rename(columns={'动态权益': '总资产'})

    fund_today2 = fund_df2[fund_df2['交易日'] == today].loc[:, ['交易日', '动态权益', '保证金']] \
        .assign(账户='国泰君安')\
        .assign(初始资产=3000000)\
        .rename(columns={'动态权益': '总资产'})

    fund_total = pd.concat([fund_today1, fund_today2]).set_index(['账户'])
    fund_total.loc['合计'] = fund_total.sum(axis=0)
    fund_total['交易日'] = today
    fund_total['资金使用率'] = fund_total['保证金'] / fund_total['总资产']
    fund_total['收益率'] = fund_total['总资产'] / fund_total['初始资产'] - 1

    fund_total['初始资产'] = fund_total['初始资产'].apply(lambda x: '%.0f' % x)
    fund_total['总资产'] = fund_total['总资产'].apply(lambda x: '%.0f' % x)
    fund_total['保证金'] = fund_total['保证金'].apply(lambda x: '%.0f' % x)
    fund_total['资金使用率'] = fund_total['资金使用率'].apply(lambda x: '%.2f%%' % (x * 100))
    fund_total['收益率'] = fund_total['收益率'].apply(lambda x: '%.2f%%' % (x * 100))
    fund_total = fund_total.loc[:, ['初始资产', '总资产', '保证金', '收益率', '资金使用率']].reset_index(drop=False)
    fund_total = fund_total.loc[:, ['账户', '初始资产', '总资产', '保证金', '收益率', '资金使用率']]

    save_df_to_doc(document, fund_total, '账户总览')

    hold_df1 = pd.read_excel(path + 'hold_' + account1 + '_' + today + '.xlsx') \
        .assign(交易日=lambda df: df['交易日'].apply(lambda x: str(int(x + 0.1))))
    hold_df1 = hold_df1[hold_df1['总仓'] != 0]
    hold_df2 = pd.read_excel(path + 'hold_' + account2 + '_' + today + '.xlsx') \
        .assign(交易日=lambda df: df['交易日'].apply(lambda x: str(int(x + 0.1))))
    hold_df2 = hold_df2[hold_df2['总仓'] != 0]
    hold1 = hold_df1.loc[:, ['合约', '方向', '总仓', '保证金', '开仓均价', '本次结算价', '资金占比', '帐号']]
    hold1['trend'] = hold1['方向'].apply(lambda x: 1 if x == '多' else -1)
    hold1['盈利价差'] = (hold1['本次结算价'] - hold1['开仓均价']) * hold1['trend']
    hold1 = hold1.loc[:, ['合约', '方向', '总仓', '保证金', '开仓均价', '盈利价差', '资金占比', '帐号']]

    hold2 = hold_df2.loc[:, ['合约', '方向', '总仓', '保证金', '开仓均价', '本次结算价', '资金占比', '帐号']]
    hold2['trend'] = hold2['方向'].apply(lambda x: 1 if x == '多' else -1)
    hold2['盈利价差'] = (hold2['本次结算价'] - hold2['开仓均价']) * hold2['trend']
    hold2 = hold2.loc[:, ['合约', '方向', '总仓', '保证金', '开仓均价', '盈利价差', '资金占比', '帐号']]

    hold1['盈利价差'] = hold1['盈利价差'].apply(lambda x: '%.0f' % x)
    hold1['开仓均价'] = hold1['开仓均价'].apply(lambda x: '%.0f' % x)
    hold1['保证金'] = hold1['保证金'].apply(lambda x: '%.0f' % x)
    hold1['资金占比'] = hold1['资金占比'].apply(lambda x: '%.2f%%' % (x * 100))

    hold2['盈利价差'] = hold2['盈利价差'].apply(lambda x: '%.0f' % x)
    hold2['开仓均价'] = hold2['开仓均价'].apply(lambda x: '%.0f' % x)
    hold2['保证金'] = hold2['保证金'].apply(lambda x: '%.0f' % x)
    hold2['资金占比'] = hold2['资金占比'].apply(lambda x: '%.2f%%' % (x * 100))

    temp = pd.DataFrame()
    for stock_name in hold_df2['合约']:
        print(stock_name)
        stock_code = stock_name
        code_df2 = hold_df2[hold_df2['合约'] == stock_name]
        stock_num2 = code_df2['总仓'].iloc[0]
        stock_cost2 = code_df2['开仓均价'].iloc[0]

        # 两个账户同时持有一只股票
        if stock_name in hold_df1['合约'].to_list():
            code_df1 = hold_df1[hold_df1['合约'] == stock_name]

            stock_num1 = code_df1['总仓'].iloc[0]
            stock_cost1 = code_df1['开仓均价'].iloc[0]
            trend = code_df1['方向'].iloc[0]
            stock_num = stock_num1 + stock_num2  # 持股数
            stock_cost = ((stock_cost1 * stock_num1) + (stock_cost2 * stock_num2)) / stock_num  # 平均成本

            if trend == '多':
                tp_point = code_df1['本次结算价'].iloc[0] - stock_cost
            else:
                tp_point = -code_df1['本次结算价'].iloc[0] + stock_cost
            margin = code_df1['保证金'].iloc[0] + code_df2['保证金'].iloc[0]

        else:
            trend = code_df2['方向'].iloc[0]
            stock_num = stock_num2  # 持股数
            stock_cost = stock_cost2  # 平均成本

            if trend == '多':
                tp_point = code_df2['本次结算价'].iloc[0] - stock_cost
            else:
                tp_point = -code_df2['本次结算价'].iloc[0] + stock_cost
            margin = code_df2['保证金'].iloc[0]

        temp = temp.append(
            {'合约': stock_name, '方向': trend, '总仓': stock_num, '开仓均价': stock_cost, '盈利价差': tp_point,
             '保证金': margin}, ignore_index=True)

    # 处理df
    for stock_name in hold_df1['合约']:
        if stock_name not in hold_df2['合约'].to_list():
            print(stock_name)
            stock_code = stock_name
            code_df1 = hold_df1[hold_df1['合约'] == stock_name]
            stock_num1 = code_df1['总仓'].iloc[0]
            stock_cost1 = code_df1['开仓均价'].iloc[0]
            trend = code_df1['方向'].iloc[0]
            stock_num = stock_num1  # 持股数
            stock_cost = stock_cost1  # 平均成本
            if trend == '多':
                tp_point = code_df1['本次结算价'].iloc[0] - stock_cost
            else:
                tp_point = -code_df1['本次结算价'].iloc[0] + stock_cost
            margin = code_df1['保证金'].iloc[0]
            temp = temp.append(
                {'合约': stock_name, '方向': trend, '总仓': stock_num, '开仓均价': stock_cost, '盈利价差': tp_point,
                 '保证金': margin}, ignore_index=True)
    print(fund_total['总资产'].iloc[-1])
    temp['仓位占比'] = temp['保证金'] / int(fund_total['总资产'].iloc[-1])

    temp['保证金'] = temp['保证金'].apply(lambda x: '%.0f' % x)
    temp['开仓均价'] = temp['开仓均价'].apply(lambda x: '%.0f' % x)
    temp['盈利价差'] = temp['盈利价差'].apply(lambda x: '%.0f' % x)
    temp['仓位占比'] = temp['仓位占比'].apply(lambda x: '%.2f%%' % (x * 100))
    temp['总仓'] = temp['总仓'].apply(lambda x: '%.0f' % x)
    temp = temp.loc[:, ['合约', '方向', '总仓', '开仓均价', '盈利价差', '保证金', '仓位占比']]

    trad_df1 = pd.read_excel(path + 'trade_' + account1 + '_' + today + '.xlsx')\
                   .loc[:, ['合约', '开平', '方向', '成交价', '成交手数', '成交时间', '成交编号']]
    # trad_df1['交易原因'] = ''
    trad_df2 = pd.read_excel(path + 'trade_' + account2 + '_' + today + '.xlsx')\
                   .loc[:, ['合约', '开平', '方向', '成交价', '成交手数', '成交时间', '成交编号']]
    # trad_df2['交易原因'] = ''


    save_df_to_doc(document, temp, '总持仓汇总')
    # save_df_to_doc(document, pd.DataFrame(), '持仓品种重大事项提醒')
    # save_df_to_doc(document, pd.DataFrame(), '交易计划')
    document.add_heading(f'浙商期货账户')
    # save_df_to_doc(document, pd.DataFrame(), '浙商期货')
    save_df_to_doc(document, pd.DataFrame(), '净值曲线')
    document.add_picture(f'{path}/fig/{account1}_net_{today}.png', width=Inches(6.0))
    save_df_to_doc(document, hold1, '期末持仓')
    save_df_to_doc(document, trad_df1, '当日交易')
    # save_df_to_doc(document, pd.DataFrame(), '国泰君安')
    document.add_heading(f'国泰君安账户')
    save_df_to_doc(document, pd.DataFrame(), '净值曲线')
    document.add_picture(f'{path}/fig/{account2}_net_{today}.png', width=Inches(6.0))
    save_df_to_doc(document, hold2, '期末持仓')
    save_df_to_doc(document, trad_df2, '当日交易')




    document.save(f'{path}/期货交易报告{today}.docx')

    a = 0






    # with MongoHandle.mongo_connect() as cl:
    #     fund_df = pd.DataFrame(cl['hosts'][host_name + '_fund'].find({'交易日': date}, {'_id': 0}))
    #     fund_df = fund_df[['策略名称', '交易日', '资产总值', '证券市值', '可用余额', '总盈亏', '净值份额', '净值']]
    #     fund_df[['资产总值', '证券市值', '可用余额', '总盈亏']] = \
    #         fund_df[['资产总值', '证券市值', '可用余额', '总盈亏']].round(2)
    #     fund_df['净值'] = fund_df['净值'].round(4)
    #     save_df_to_doc(document, fund_df, '策略概要')
    #     deal_df = pd.DataFrame(cl['hosts'][host_name + '_deal'].find({'交易日': date}, {'_id': 0}))
    #     if not deal_df.empty:
    #         deal_df = deal_df[['证券代码', '证券名称', '交易类型', '成交价格', '成交数量', '成交时间',
    #                            '成交金额', '手续费', '策略名称']]
    #         deal_df['手续费'] = deal_df['手续费'].round(2)
    #     else:
    #         deal_df = pd.DataFrame(columns=['证券代码', '证券名称', '交易类型',
    #                                         '成交价格', '成交数量', '成交时间', '成交金额', '手续费', '策略名称'])
    #     save_df_to_doc(document, deal_df, '交易统计')
    #     hold_df = pd.DataFrame(cl['hosts'][host_name + '_hold'].find({'交易日': date}, {'_id': 0}))
    #     hold_df = hold_df[['证券代码', '证券名称', '当前拥股数量', '冻结数量', '可卖数量', '成本价格', '市场价',
    #                        '浮动盈亏', '盈亏比例', '证券市值']]
    #     hold_df.rename(columns={'当前拥股数量': '余股', '冻结数量': '冻结', '可卖数量': '可卖'}, inplace=True)
    #     hold_df['成本价格'] = hold_df['成本价格'].round(2)
    #     hold_df['浮动盈亏'] = hold_df['浮动盈亏'].round(2)
    #     save_df_to_doc(document, hold_df, '期末持仓')
    #     # document.add_paragraph(u'/n')
    #     # document.add_picture('E:/OneDrive/Workspace/股票/交易报告/预期策略模拟盘交易报告/1.png', width=Inches(6.0))
    #     # document.add_picture('E:/OneDrive/Workspace/股票/交易报告/预期策略模拟盘交易报告/2.png', width=Inches(6.0))
    #     # document.add_picture('E:/OneDrive/Workspace/股票/交易报告/预期策略模拟盘交易报告/3.png', width=Inches(6.0))
    #     document.save(f'{path}/{host_name}策略交易报告{date}.docx')
    #     print('报告生成完成')